import camelot
import matplotlib.pyplot as plt
import pandas as pd
import re 
#from barcode import Code128
#from barcode.writer import ImageWriter
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import RGBColor, Pt
from pypdf import PdfReader, PdfWriter, Transformation, PageObject, PaperSize
from pypdf.generic import RectangleObject

from code128 import format

#Create an instance of docx to edit
document = Document()
#Assign the document's last one section (all pages)
section = document.sections[-1]
#Set orientation to landscape for whole document
section.orientation = WD_ORIENT.LANDSCAPE
#Manually set the page width and height after setting landscape, else it won't adjust to landscape
new_width, new_height = section.page_height, section.page_width
section.page_width = new_width
section.page_height = new_height

#The PDF file for the ends plan
file = 'plan.pdf'

#Open Pdf to get dimensions on main bargen() call
reader = PdfReader(file)

#Open instance of PdfWriter to create new pdf for resized pages
writer = PdfWriter()

#Base page dimensions to resize oversized pages
resize_h = 612
resize_w = 792

#Resizes target PDF page and adds to writer instance
def pdf_resizer (page_index):
    #Set box of the target page
    box = reader.pages[page_index]
    #Set the scaling to downscale
    scale_factor = min(resize_h/box.mediabox.height, resize_w/box.mediabox.width)
    #Create transformation instance with that scale
    transform = Transformation().scale(scale_factor, scale_factor)
    #Apply the transformation on the box
    box.add_transformation(transform)
    #Crop out the rest of the page
    box.cropbox = RectangleObject((0,0,resize_w, resize_h))
    #Create blank page
    resized_page = PageObject.create_blank_page(width = resize_w, height = resize_h)
    #Set box size same as blank page size
    box.mediabox = resized_page.mediabox
    #Merge the cropped box onto the created blank page
    resized_page.merge_page(box)
    #Add the created page into the writer instance
    writer.add_page(resized_page)

pdf_resizer(0)
pdf_resizer(1)

#Name of resized PDF file and create the PDF to local folder.
resized_file = 'resized.pdf'
writer.write(resized_file)

#Used to rescale the pdf regions, table_areas instead of table_regions to measure an area
#data = camelot.read_pdf(resized_file, pages='1', flavor='lattice', table_areas=['5, 185, 235, 5']) #Ends 9
#data = camelot.read_pdf(resized_file, pages='1', flavor='lattice', table_regions=['5, 330, 235, 160']) #Ends 5
#data = camelot.read_pdf(resized_file, pages='1', flavor='lattice',table_areas=['5, 470, 235, 300']) #Ends 1
#data = camelot.read_pdf(resized_file, pages='1', flavor='lattice', table_areas=['200, 185, 405, 5']) #Ends 10
#data = camelot.read_pdf(resized_file, pages='1', flavor='lattice', table_regions=['200, 330, 405, 160']) #Ends 6
#data = camelot.read_pdf(resized_file, pages='1', flavor='lattice', table_areas=['200, 470, 405, 300']) #Ends 2
#data = camelot.read_pdf(resized_file, pages='1', flavor='lattice', table_areas=['390, 185, 600, 5']) #Ends 11
#data = camelot.read_pdf(resized_file, pages='1', flavor='lattice', table_areas=['390, 330, 600, 160']) #Ends 7
#data = camelot.read_pdf(resized_file, pages='1', flavor='lattice', table_areas=['390, 470, 600, 300']) #Ends 3
#data = camelot.read_pdf(resized_file, pages='1', flavor='lattice', table_areas=['570, 185, 790, 5']) #Ends 12
#data = camelot.read_pdf(resized_file, pages='1', flavor='lattice', table_areas=['570, 330, 790, 160']) #Ends 8
#data = camelot.read_pdf(resized_file, pages='1', flavor='lattice', table_areas=['570, 470, 790, 300']) #Ends 4

#data = camelot.read_pdf(resized_file, pages='2', flavor='lattice', table_areas=['5, 205, 235, 5']) #Ends 9
#data = camelot.read_pdf(resized_file, pages='2', flavor='lattice', table_areas=['5, 360, 235, 180']) #Ends 5
#data = camelot.read_pdf(resized_file, pages='2', flavor='lattice', table_areas=['5, 530, 235, 350']) #Ends 1
#data = camelot.read_pdf(resized_file, pages='2', flavor='lattice', table_areas=['200, 205, 405, 5']) #Ends 10
#data = camelot.read_pdf(resized_file, pages='2', flavor='lattice', table_areas=['200, 360, 405, 180']) #Ends 6
#data = camelot.read_pdf(resized_file, pages='2', flavor='lattice', table_areas=['200, 530, 405, 350']) #Ends 2
#data = camelot.read_pdf(resized_file, pages='2', flavor='lattice', table_areas=['390, 205, 600, 5']) #Ends 11
#data = camelot.read_pdf(resized_file, pages='2', flavor='lattice', table_areas=['390, 360, 600, 180']) #Ends 7
#data = camelot.read_pdf(resized_file, pages='2', flavor='lattice', table_areas=['390, 530, 600, 350']) #Ends 3
#data = camelot.read_pdf(resized_file, pages='2', flavor='lattice', table_areas=['570, 205, 790, 5']) #Ends 12
#data = camelot.read_pdf(resized_file, pages='2', flavor='lattice', table_areas=['570, 360, 790, 180']) #Ends 8
#data = camelot.read_pdf(resized_file, pages='2', flavor='lattice', table_areas=['570, 530, 790, 350']) #Ends 4

#Regions (x1, y1, x2, y2) top left and bottom right respectively of the square, to search table in for page 1 and 2.
pdf_regions_1 = [['5, 470, 235, 300'], ['200, 470, 405, 300'], ['390, 470, 600, 300'], ['570, 470, 790, 300'], ['5, 330, 235, 160'], 
                 ['200, 330, 405, 160'], ['390, 330, 600, 160'], ['570, 330, 790, 160'], ['5, 185, 235, 5'], ['200, 185, 405, 5'], ['390, 185, 600, 5'], ['570, 185, 790, 5']]
pdf_regions_2 = [['5, 530, 235, 350'], ['200, 530, 405, 350'], ['390, 530, 600, 350'], ['570, 530, 790, 350'], ['5, 360, 235, 180'], 
                 ['200, 360, 405, 180'], ['390, 360, 600, 180'], ['570, 360, 790, 180'], ['5, 205, 235, 5'], ['200, 205, 405, 5'], ['390, 205, 600, 5'], ['570, 205, 790, 5']]

#Plot the area on the PDF
#camelot.plot(data[0],kind='contour')
#plt.show(block=True)

#Function to edit docx instance with data from the dataframes, modifying the data, and generating barcodes.
def generate_doc(dataframe):
    #Counters for incrementing shelf numbers and their names
    shelf_counter = 1
    #name_counter = 0
    for column in dataframe:
        for row in dataframe[column]:
            #Check each row for front ends and back ends to assign to a new page
            if re.search("FGE|OGE", row):
                shelf_counter = 1
                document.add_page_break()
                paragraph = document.add_paragraph()
                run = paragraph.add_run(row + "\n")
                run.bold = True
                #If the end is other than shelves, change font to red
                if (re.search("pre-pk|pallet|tray|bin", row, re.IGNORECASE)):
                    run.font.color.rgb = RGBColor(255, 0, 0)
            else:
                #Add commas for references with spaces inbetween
                add_commas = re.sub("(?<=\d)(\s+)(?=\d)",",", row)
                #Trim out all white spaces
                rmv_whitespaces = re.sub("\s+", "", add_commas)
                #Substitute "or" with a comma between references
                replace_or = re.sub("(?<=\d)(or)(?=\d)",",", rmv_whitespaces)
                #Replace all plus signs with a comma
                content = replace_or.replace("+",",")
                #Match numbers after "Ref:"
                matched_codes = re.findall("Ref:(\d+(?:,\d+)*\d)", content)
                #Match names before "Ref:"
                matched_names = re.findall("^.+?(?=Ref:)", content)
                #Check for scenarios where row has no "Ref:" (has code, but no name for it), and when there's a name, but no Ref code.
                if matched_codes and matched_names == []:
                    matched_names = ['CHECK SOURCE MATERIAL']
                elif matched_names and matched_codes == []:
                    document.add_paragraph("CHECK SOURCE MATERIAL")                 
                if matched_codes:
                    #Write a line for shelf number and shelf name, increment the shelf number for the next.
                    document.add_paragraph("Shelf " + str(shelf_counter) + " " + matched_names[0])
                    shelf_counter += 1
                    for codes in matched_codes:
                        #Create new array of singular codes to generate each barcode
                        code_array = codes.split(',')
                        #Add another paragraph
                        paragraph = document.add_paragraph()
                        #Generate each barcode as PNG and adding to the docx instance, replacing the same image file locally.
                        for code in code_array:

                            # barcode = Code128(code, writer = ImageWriter())
                            # barcode.save("barcode", {"module_width": 0.15, "module_height": 3, "font_size": 2.5, "text_distance": 1.5, "quiet_zone": 2})
                            # run.add_picture("barcode.png")

                            # Generate ASCII character string from the adjusted Code128 package format.py file
                            ascii_char_string = format.code128_format(code)
                            print(ascii_char_string)
                            #Append first run to this paragraph allowing barcode above human code with styling
                            barcode_run = paragraph.add_run(ascii_char_string)
                            font = barcode_run.font
                            font.name = "Libre Barcode 128"
                            font.size = Pt(48)
                            #Append second run to this paragraph allowing human readable code
                            paragraph.add_run("\n\t" + code + "\n")

                            
                            
#Reads that PDF file to create dataframes and edit the document instance for each scaled PDF region
def scrape_pdf(page_num, pdf_regions):
    #If page is A4 size or less, scrape regions of the page for tables, otherwise scrape whole page. Index out of range error means table wasn't found. 
    try:
        for i in range(len(pdf_regions)):
            data = camelot.read_pdf(resized_file, pages=page_num, flavor='lattice', table_regions=pdf_regions[i])
            df = data[0].df
            generate_doc(df)
            print(i)
    except IndexError as e:
        data = camelot.read_pdf(resized_file, pages=page_num, flavor='lattice')
        df = data[0].df
        generate_doc(df)
        print("A table wasn't found - {}".format(e))

#The main function for front and back ends, generating shelf number, end names, shelf names and barcodes for the document instance.
def bargen():
    scrape_pdf('1', pdf_regions_1) #Front ends
    scrape_pdf('2', pdf_regions_2) #Back ends

#Execute main
bargen()


#Save the completed docx instance as a file in the local folder, overwriting existing.
document.save("results.docx")